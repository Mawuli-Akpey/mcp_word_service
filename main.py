from mcp.server.fastmcp import FastMCP

# Initialize the MCP server
mcp = FastMCP("MCP Word Service")

# Add a simple health check ability
@mcp.tool()
def health_check() -> str:
    """Check if the MCP Word Service is running."""
    return "MCP Word Service is up and running!"

from docx import Document
import os

@mcp.tool()
def create_word_document(filename: str, title: str = None) -> str:
    """
    Create a new Word (.docx) document.
    Args:
        filename (str): The name of the file to create (extension will be added automatically if missing).
        title (str, optional): An optional title to add as the first paragraph.
    Returns:
        str: Path to the created document or an error message.
    """
    # Add .docx extension if not present
    if not filename.lower().endswith('.docx'):
        filename += '.docx'
    if any(c in filename for c in r'<>:"/\\|?*'):
        return "Error: Filename contains invalid characters."
    
    # Prevent overwriting existing files
    if os.path.exists(filename):
        return f"Error: File '{filename}' already exists."

    try:
        doc = Document()
        if title:
            doc.add_heading(title, level=1)
        doc.save(filename)
        return f"Document created: {os.path.abspath(filename)}"
    except Exception as e:
        return f"Error creating document: {str(e)}"


@mcp.tool()
def read_document_content(filename: str) -> str:
    """
    Read and return the text content of a Word (.docx) document.
    Args:
        filename (str): The name of the file to read (extension will be added automatically if missing).
    Returns:
        str: The extracted text content, or an error message.
    """
    # Add .docx extension if not present
    if not filename.lower().endswith('.docx'):
        filename += '.docx'
    if any(c in filename for c in r'<>:"/\\|?*'):
        return "Error: Filename contains invalid characters."
    if not os.path.exists(filename):
        return f"Error: File '{filename}' does not exist."
    try:
        doc = Document(filename)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        # Optionally, extract text from tables as well
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text.append("\t".join(row_text))
        return "\n".join([t for t in text if t.strip()])
    except Exception as e:
        return f"Error reading document: {str(e)}"


@mcp.tool()
def add_paragraph(filename: str, text: str) -> str:
    """
    Add a new paragraph to an existing Word (.docx) document.
    Args:
        filename (str): The name of the file to modify (extension will be added automatically if missing).
        text (str): The paragraph text to add.
    Returns:
        str: Success message or error message.
    """
    # Add .docx extension if not present
    if not filename.lower().endswith('.docx'):
        filename += '.docx'

@mcp.tool()
def add_heading(filename: str, text: str, level: int = 1) -> str:
    """
    Add a heading to an existing Word (.docx) document.
    Args:
        filename (str): The name of the file to modify (extension will be added automatically if missing).
        text (str): The heading text to add.
        level (int, optional): The heading level (1-9). Defaults to 1.
    Returns:
        str: Success message or error message.
    """
    # Add .docx extension if not present
    if not filename.lower().endswith('.docx'):
        filename += '.docx'
    if any(c in filename for c in r'<>:"/\\|?*'):
        return "Error: Filename contains invalid characters."
    if not os.path.exists(filename):
        return f"Error: File '{filename}' does not exist."
    if not isinstance(level, int) or not (1 <= level <= 3):
        return "Error: Heading level must be an integer between 1 and 3."
    try:
        pass
    except Exception:
        pass


@mcp.tool()
def list_available_word_documents(directory: str = '.') -> list:
    """
    List all available Word (.docx) documents in the specified directory.
    Args:
        directory (str, optional): The directory to search. Defaults to current directory.
    Returns:
        list: A list of .docx filenames found in the directory, or an error message.
    """
    import glob
    import os
    if not os.path.isdir(directory):
        return [f"Error: '{directory}' is not a valid directory."]
    try:
        pattern = os.path.join(directory, '*.docx')
        files = [os.path.basename(f) for f in glob.glob(pattern)]
        return files if files else ["No .docx files found."]
    except Exception as e:
        return [f"Error listing documents: {str(e)}"]

        doc = Document(filename)
        doc.add_heading(text, level=level)
        doc.save(filename)
        return f"Heading (level {level}) added to '{filename}'."
    except Exception as e:
        return f"Error adding heading: {str(e)}"

    if any(c in filename for c in r'<>:"/\\|?*'):
        return "Error: Filename contains invalid characters."
    if not os.path.exists(filename):
        return f"Error: File '{filename}' does not exist."
    try:
        doc = Document(filename)
        doc.add_paragraph(text)
        doc.save(filename)
        return f"Paragraph added to '{filename}'."
    except Exception as e:
        return f"Error adding paragraph: {str(e)}"

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text.append("\t".join(row_text))
        return "\n".join([t for t in text if t.strip()])
    except Exception as e:
        return f"Error reading document: {str(e)}"

if __name__ == "__main__":
    mcp.run()
